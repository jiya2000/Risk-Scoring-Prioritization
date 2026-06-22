"""
Fill Invention_Disclosure_Format_B.docx with all content from the LaTeX IDF-B.
Run from the project root.
"""
import copy, os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants

SRC  = r"C:\Users\abhis\Desktop\Risk-Scoring-Prioritization-main\Risk-Scoring-Prioritization-main\Invention_Disclosure_Format_B.docx"
OUT  = r"C:\Users\abhis\Desktop\Risk-Scoring-Prioritization-main\Risk-Scoring-Prioritization-main\Invention_Disclosure_Format_B_Filled.docx"
IMG  = r"C:\Users\abhis\Desktop\Risk-Scoring-Prioritization-main\Risk-Scoring-Prioritization-main\docs\artifacts"

doc = Document(SRC)

# ──────────────────────────────────────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def bold_run(para, text, size=10):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return run

def normal_run(para, text, size=10):
    run = para.add_run(text)
    run.bold = False
    run.font.size = Pt(size)
    return run

def add_heading(doc, text, level=2, color='1F3864'):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading %d' % level] if ('Heading %d' % level) in [s.name for s in doc.styles] else doc.styles['Normal']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11 if level == 2 else 10)
    run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_body(doc, text, size=10, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_numbered(doc, text, size=10):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_image(doc, fname, width=6.0, caption_text=''):
    path = os.path.join(IMG, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption_text:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption_text)
        r.italic = True
        r.font.size = Pt(9)

def add_table_row(table, cells, bold_first=False, bg=None):
    row = table.add_row()
    for i, (cell, val) in enumerate(zip(row.cells, cells)):
        cell.text = str(val)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
            if bold_first and i == 0:
                run.bold = True
        if bg:
            set_cell_bg(cell, bg)
    return row

def make_table(doc, headers, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hrow = t.rows[0]
    for i, (cell, h) in enumerate(zip(hrow.cells, headers)):
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        set_cell_bg(cell, '1F3864')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    return t

# ──────────────────────────────────────────────────────────────────────
# Locate the 10 section header paragraphs and clear blank fillers
# ──────────────────────────────────────────────────────────────────────
# Strategy: we will INSERT content after each section heading paragraph.
# We identify each heading by its index and insert paragraphs after it.

from docx.oxml.ns import qn as _qn
from lxml import etree

def insert_after(ref_para, new_para):
    """Insert new_para's XML element immediately after ref_para in the doc body."""
    ref_para._element.addnext(new_para._element)

def para_after_index(doc, idx, text, bold=False, italic=False, size=10, indent=0):
    """Create a paragraph and insert it after paragraph at idx."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    if indent:
        ind = OxmlElement('w:ind')
        ind.set(_qn('w:left'), str(indent))
        pPr.append(ind)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz'); sz.set(_qn('w:val'), str(int(size*2)))
    szCs = OxmlElement('w:szCs'); szCs.set(_qn('w:val'), str(int(size*2)))
    rPr.append(sz); rPr.append(szCs)
    if bold:
        b = OxmlElement('w:b'); rPr.append(b)
    if italic:
        i = OxmlElement('w:i'); rPr.append(i)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); p.append(r)
    return p

# Remove ALL blank body-text paragraphs between section headings
# (keep section heading paragraphs and the END line)
to_remove = []
for p in doc.paragraphs:
    if p.style.name == 'Body Text' and p.text.strip() == '':
        to_remove.append(p._element)
    if p.style.name == 'List Paragraph' and p.text.strip() == '':
        to_remove.append(p._element)

for el in to_remove:
    el.getparent().remove(el)

# Re-read paragraphs after cleanup
paras = doc.paragraphs

print("Paragraphs after cleanup:")
for i, p in enumerate(paras):
    print(f"  [{i}] {p.text[:60]!r}")

# ──────────────────────────────────────────────────────────────────────
# Helper: append rich content block to the document body
# We'll build content in a separate document and merge elements
# ──────────────────────────────────────────────────────────────────────

body = doc.element.body   # main document body XML element

def ap(text, bold=False, italic=False, size=10, indent=False, align_center=False, color=None):
    """Append paragraph to body."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def ah(text, color='1F3864', size=11):
    """Append a bold heading paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

_numbered_counter = [0]

def ab(text, size=10):
    """Append bullet point using Normal style + bullet prefix."""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(u"\u2022  " + text)
    run.font.size = Pt(size)
    return p

def an(text, size=10):
    """Append numbered item using Normal style + auto-counter."""
    _numbered_counter[0] += 1
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.left_indent = Inches(0.35)
    run = p.add_run(f"{_numbered_counter[0]}.  {text}")
    run.font.size = Pt(size)
    return p

def reset_counter():
    _numbered_counter[0] = 0

def aimg(fname, width=6.2, cap=''):
    path = os.path.join(IMG, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(cap)
        r.italic = True
        r.font.size = Pt(9)

def sep():
    ap('')

def mk_table(headers, rows, col_widths_cm=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Normal'
    # Add borders manually via XML
    from docx.oxml import OxmlElement as OE
    from docx.oxml.ns import qn as QN
    def set_borders(tbl):
        tblPr = tbl._tbl.tblPr
        if tblPr is None:
            tblPr = OE('w:tblPr')
            tbl._tbl.insert(0, tblPr)
        tblBorders = OE('w:tblBorders')
        for border in ['top','left','bottom','right','insideH','insideV']:
            el = OE(f'w:{border}')
            el.set(QN('w:val'),   'single')
            el.set(QN('w:sz'),    '4')
            el.set(QN('w:space'), '0')
            el.set(QN('w:color'), '444444')
            tblBorders.append(el)
        tblPr.append(tblBorders)
    set_borders(t)
    # header row
    hr = t.rows[0]
    for i, (c, h) in enumerate(zip(hr.cells, headers)):
        c.text = ''
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(c, '1F3864')
    # data rows
    for row_data in rows:
        dr = t.add_row()
        for i, (c, val) in enumerate(zip(dr.cells, row_data)):
            c.text = ''
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
            if i == 0:
                run.bold = True
    return t

# ══════════════════════════════════════════════════════════════════════
# Move the END paragraph to the very bottom, insert content before it
# We'll simply append everything and then re-append the END line.
# ══════════════════════════════════════════════════════════════════════

# Find section heading paras (after cleanup)
paras = doc.paragraphs
sec = {}
for i, p in enumerate(paras):
    t = p.text.strip()
    if 'Title of the invention'           in t: sec[1]  = p
    if 'Field /Area of invention'         in t: sec[2]  = p
    if 'Prior Patents and Publications'   in t: sec[3]  = p
    if 'Summary and background'           in t: sec[4]  = p
    if 'Objective(s) of Invention'        in t: sec[5]  = p
    if 'Working principle'                in t: sec[6]  = p
    if 'Description of the invention in detail' in t: sec[7] = p
    if 'Experimental validation'          in t: sec[8]  = p
    if 'What aspect(s) of the invention'  in t: sec[9]  = p
    if 'Technology readiness level'       in t: sec[10] = p

print("Found sections:", list(sec.keys()))

# ── We'll BUILD new paragraphs/tables and insert them after each heading ──
# Using document.add_* appends to the END, then we relocate. 
# Simpler: move END marker first, then append all content, then add END back.

# Save and remove the END paragraph
end_para_el = None
for p in doc.paragraphs:
    if 'END OF THE DOCUMENT' in p.text:
        end_para_el = p._element
        p._element.getparent().remove(p._element)
        break

# Also remove the section heading paragraphs themselves – 
# we will re-add them with content below from scratch in proper order.
# Actually keep them but update their text to be numbered properly.

# Rename/update each heading text to be numbered
heading_texts = {
    1:  "1.  Title of the Invention",
    2:  "2.  Field / Area of Invention",
    3:  "3.  Prior Patents and Publications from Literature",
    4:  "4.  Summary and Background of the Invention (Address the Gap / Novelty)",
    5:  "5.  Objective(s) of Invention",
    6:  "6.  Working Principle of the Invention (In Brief)",
    7:  "7.  Description of the Invention in Detail",
    8:  "8.  Experimental Validation Results",
    9:  "9.  What Aspect(s) of the Invention Need(s) Protection?",
    10: "10. Technology Readiness Level (TRL)",
}
for k, p in sec.items():
    for run in p.runs:
        run.text = ''
    run = p.add_run(heading_texts[k])
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

# ══════════════════════════════════════════════════════════════════════
# SECTION 1 — TITLE
# ══════════════════════════════════════════════════════════════════════
ap('')
ap("Topology-Conditioned Adaptive Score Fusion with Temporal-Decay Graph Centrality, "
   "Precision-Guaranteed Degradation Control, and Learnable Architecture Hardening for "
   "Anti-Money Laundering Risk Prioritization",
   bold=True, size=11)
sep()
ap("Short form:", bold=True, size=10)
ap("AML Risk Scoring Platform with Patentable Innovations: TD-PageRank, TopologyAttentionGate "
   "Fusion, Adaptive Precision-Budget Degradation Controller, Learnable SCC Penalty, "
   "GNN Topology Embedding, Online Precision Recalibration, and System-Level §101 "
   "Resource Management Anchoring", italic=True, size=10)
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 2 — FIELD
# ══════════════════════════════════════════════════════════════════════
ap('')
ab("Primary Field: Artificial Intelligence and Machine Learning — graph-based financial anomaly detection and ensemble risk scoring")
ab("Sub-field 1: Graph Algorithm Design — temporal-decay modification of the PageRank power iteration for financial transaction networks")
ab("Sub-field 2: Neural Architecture — multiplicative attention gating, GNN-based ego-network embedding, learned isolation detection for heterogeneous signal fusion")
ab("Sub-field 3: System Reliability Engineering — precision-budget-constrained adaptive degradation control with online precision recalibration")
ab("Sub-field 4: Patent §101 System-Level Resource Management — concrete execution-path switching, memory buffer allocation, atomic model swaps tied to algorithmic innovations")
ab("Application Domain: Anti-Money Laundering (AML) Compliance, Financial Crime Detection, Bank Transaction Risk Prioritisation")
ab("Regulatory Framework: FATF Recommendation 16; Nepal Rastra Bank AML Directives 2023; FinCEN SAR Filing Requirements (31 CFR 1020.320)")
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 3 — PRIOR ART TABLE
# ══════════════════════════════════════════════════════════════════════
ap('')
mk_table(
    ["#", "Reference", "Type", "Year", "What It Does", "Gap Addressed"],
    [
        ["1", "US20220405860A1 — Fraud Detection Using Weighted Ensemble of Heterogeneous Models", "Patent", "2022",
         "Combines ML, graph, and rule scores using fixed global weights learned at training time",
         "Does not adapt weights based on local graph topology of each individual account"],
        ["2", "US11640609B1 — Network-Based Features for Financial Crime (Wells Fargo)", "Patent", "2023",
         "Risk-vector propagation x=A^p·z on discretely time-binned adjacency matrices",
         "No exponential temporal decay inside iteration; no burst-velocity amplification; no directional SCC penalty"],
        ["3", "US20210174258A1 — Machine Learning Monitoring Systems (Microsoft)", "Patent", "2021",
         "Monitors ML pipeline health via heartbeat/latency thresholds and emits alerts",
         "Does not guarantee output Precision@K; no routing table; no self-calibrating precision budget"],
        ["4", "US10713722B2 — Graph Neural Networks for Financial Fraud Detection", "Patent", "2020",
         "GNN-based node embeddings for fraud node classification",
         "Not applied to ego-network topology for per-account fusion weight prediction; no topology-conditioned gating"],
        ["5", "Alice Corp. v. CLS Bank, 573 U.S. 208 (2014)", "Legal", "2014",
         "Abstract ideas not patentable under §101 without 'something more'",
         "Addressed by our ResourceManager: concrete memory buffer allocation, execution-path routing, atomic model swaps"],
        ["6", "US20230267516A1 — Adaptive Model Monitoring", "Patent", "2023",
         "Monitors ML model drift with static alarm thresholds",
         "No adaptive precision budget; no per-path precision routing table; no online Precision@50 estimation"],
        ["7", "Langville & Meyer (2006) — Google's PageRank and Beyond", "Pub.", "2006",
         "Foundational PageRank with uniform edge weights and teleportation",
         "No temporal decay, no transaction-amount weighting, no AML pattern awareness"],
        ["8", "Chen et al. (2020) — Graph Attention Networks for Financial Fraud", "Pub.", "2020",
         "Standard additive attention on feature space for fraud classification",
         "Not multiplicative gating on the weight dimension of heterogeneous signal ensembles"],
        ["9", "Netflix Hystrix / Microsoft Polly — Circuit Breaker Pattern", "OSS", "2012",
         "HEALTHY/DEGRADED state machine using latency/availability as routing metrics",
         "Routing metric is latency, not Precision@K; no pre-evaluated precision routing table; fixed thresholds"],
        ["10", "FATF (2022) — International Standards on Combating Money Laundering", "Standard", "2022",
         "Defines 10+ money-laundering typologies (Smurfing, Layering, Fan-Out, etc.)",
         "Does not specify computational graph-structure detection algorithms"],
    ]
)
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 4 — SUMMARY AND BACKGROUND
# ══════════════════════════════════════════════════════════════════════
ap('')
ah("4.1  Background", size=10)
ap("AML compliance teams at financial institutions receive between 50,000 and 200,000 "
   "transaction alerts per week. Legacy rule-based alert systems produce >95% false positive "
   "rates, resulting in 'alert fatigue' — investigators waste 90% of their time on clean "
   "accounts while genuine money laundering cases accumulate undetected. Three fundamental "
   "architectural gaps exist in prior ML-based AML systems:", size=10)
sep()
ap("Gap 1 — Static ensemble weights:", bold=True, indent=True, size=10)
ap("All prior ensemble methods (including US20220405860) use globally-learned or fixed weights "
   "applied uniformly to all accounts. A smurfing-ring account embedded in a dense criminal "
   "network and an isolated account both receive identical weights — structurally incorrect.", indent=True, size=10)
sep()
ap("Gap 2 — Temporally unaware graph centrality:", bold=True, indent=True, size=10)
ap("Prior graph-based AML methods (including US11640609B1) compute centrality on discretely "
   "time-binned transaction graphs using fixed-step matrix propagation (x = A^p·z). Temporal "
   "decay is not embedded continuously in the iteration. Burst-pattern transactions — a hallmark "
   "of smurfing — receive the same weight as ordinary patterns. Cyclic fund transfers are penalized "
   "uniformly regardless of whether the node is a fund collector or distributor.", indent=True, size=10)
sep()
ap("Gap 3 — No precision guarantee during component failures:", bold=True, indent=True, size=10)
ap("Production AML systems have no mechanism to guarantee that the fraud prioritization queue "
   "maintains a minimum quality level when individual components fail. Prior art (US20210174258A1) "
   "monitors component health and emits alerts, but does not route computation through alternative "
   "execution paths with pre-validated precision scores.", indent=True, size=10)
sep()
ap("Gap 4 (New) — Static heuristics and manual features:", bold=True, indent=True, size=10)
ap("Existing SCC penalties rely on three hard-coded archetypes. Topology representation is manually "
   "engineered. Precision monitoring relies on offline-computed P@50 assumptions that become stale "
   "under concept drift.", indent=True, size=10)
sep()

ah("4.2  Summary of Three Core Innovations", size=10)
ap("Innovation 1 — Temporal-Decay PageRank (TD-PageRank):", bold=True, size=10)
ap("A modified PageRank where exponential temporal decay lambda=ln(2)/half_life_days is embedded "
   "directly inside the power iteration formula. Additionally: (a) burst-velocity amplification "
   "multiplies edge weights for sender nodes with >5 transactions in 3 days by (1+burst_ratio)>=1.0; "
   "(b) log-amount normalization prevents large single transactions from dominating centrality; "
   "(c) directional SCC penalty assigns asymmetric dampening: collectors 0.25x, distributors 0.50x, "
   "balanced 0.375x; (d) dormant node suppression caps scores at 0.1x maximum when all incident "
   "edges exceed 30 days.", indent=True, size=10)
sep()
ap("Innovation 2 — Topology-Adaptive Fusion with TopologyAttentionGate:", bold=True, size=10)
ap("A per-account dynamic weight computation system where a 6-dimensional topology vector "
   "(edge density, diameter, clustering, degree asymmetry, component ratio, velocity ratio) "
   "is extracted from each account's 2-hop ego-network and passed through a TopologyAttentionGate — "
   "a multiplicative gating neural network. This is distinct from standard additive attention: "
   "the gate amplifies the graph-signal weight for topologically active accounts and suppresses it "
   "for isolated ones, operating on the weight dimension of the ensemble.", indent=True, size=10)
sep()
ap("Innovation 3 — Adaptive Degradation Controller with Precision Budget:", bold=True, size=10)
ap("A system that monitors 5 pipeline components and routes scoring computation through one of 8 "
   "pre-evaluated execution paths, each with an empirically measured Precision@50. Novel "
   "sub-mechanisms: (a) PrecisionDriftDetector — EMA-smoothed KL-divergence monitor; "
   "(b) AdaptivePrecisionBudget — self-calibrating minimum precision threshold that tightens "
   "when consistently high (+0.05) and relaxes when borderline (-0.02), with configurable floor [0.55, 0.75].", indent=True, size=10)
sep()

ah("4.3  Architecture Hardening Innovations (New — June 2026)", size=10)
ap("All four hardening surfaces are guarded by ArchitectureHardeningConfig flags defaulting to "
   "False — fully backward-compatible with the existing baseline system.", size=10)
sep()
aimg("fig07_architecture_hardening.png", width=6.2,
     cap="Fig 7 — Architecture Hardening Layer: 4 surfaces, all flags default=False for backward compatibility")
sep()
ap("Surface 1 — Learnable SCC Penalty:", bold=True, indent=True, size=10)
ap("LearnableSCCPenalty MLP (input: 5 intra-SCC flow features -> penalty in [0.1, 1.0]) replaces "
   "the 3-archetype static heuristic. Retrained via train_penalty() with temporal splits. "
   "Updated atomically via ResourceManager.atomic_model_swap().", indent=True, size=10)
sep()
ap("Surface 2 — GNN Topology Embedding:", bold=True, indent=True, size=10)
ap("TopologyEmbeddingNetwork (2x GCNConv layers, 16-dimensional embedding) + IsolationDetector "
   "(learned continuous isolation score, monotonically non-decreasing w_ml, replaces hard-coded "
   "w_ml>=0.70 threshold) + DeepTopologyAttentionGate (>=2 hidden layers, >=64 hidden dim, "
   "dropout regularization).", indent=True, size=10)
sep()
ap("Surface 3 — Online Precision Recalibration:", bold=True, indent=True, size=10)
ap("OnlinePrecisionMonitor (rolling 500-sample window, P@50 recomputed live) + "
   "EnhancedConceptDriftDetector (KL-divergence + precision jointly classify: "
   "no_drift / benign_shift / precision_degraded, per-path isolation) + "
   "EnhancedAdaptivePrecisionBudget (0.7 x online + 0.3 x shadow blending, output in [0.55, 0.75]).", indent=True, size=10)
sep()
ap("Surface 4 — §101 Resource Management Anchoring:", bold=True, indent=True, size=10)
ap("ResourceManager records every execution-path switch with concrete resource metrics: memory "
   "allocation delta, computation-time delta, and active processing-unit count — anchoring patent "
   "claims to system-level technological improvements that survive Alice Corp. v. CLS Bank.", indent=True, size=10)
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 5 — OBJECTIVES
# ══════════════════════════════════════════════════════════════════════
ap('')
ah("Original Objectives:", size=10)
reset_counter()
an("Develop a graph centrality algorithm embedding exponential temporal decay directly in the "
   "PageRank power iteration and detecting smurfing via burst-velocity amplification, achieving "
   ">=15% mean absolute score difference from standard PageRank.")
an("Develop a per-account topology-conditioned ensemble fusion mechanism using multiplicative "
   "attention gating, achieving >=10% relative P@50 improvement over static fusion baselines.")
an("Develop a precision-budget-constrained degradation controller guaranteeing P@50 >= 0.60 "
   "across all 8 operational modes including partial component failures.")
an("Create a complete production-deployable AML platform achieving 82% P@50 on temporal test "
   "data with zero future data leakage, integrating a 10-typology symbolic rule engine, entity-"
   "faithful NLP narrative generation, and 6-level graceful degradation.")
an("Provide formal machine-verified correctness guarantees via property-based testing: 27 "
   "correctness properties, 119 tests, all proven via Hypothesis framework with 50-100 "
   "auto-generated examples each.")
sep()
ah("New Objectives (Architecture Hardening):", size=10)
reset_counter()
an("Learn optimal SCC penalty multipliers from transaction data via a trainable MLP, adapting "
   "to evolving laundering topologies beyond three hard-coded archetypes.")
an("Replace manual topology feature engineering with end-to-end learned GNN embeddings "
   "(TopologyEmbeddingNetwork) and replace hard-coded isolation thresholds with a learned "
   "continuous isolation function (IsolationDetector).")
an("Replace static offline P@50 assumptions in the routing table with live online precision "
   "recalibration (OnlinePrecisionMonitor) and direct precision drift detection "
   "(EnhancedConceptDriftDetector).")
an("Anchor all patent claims to concrete system-level resource management operations (memory "
   "buffer allocation, execution-path routing, atomic model swap, utilisation reporting) to "
   "satisfy 35 U.S.C. §101 subject-matter eligibility.")
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 6 — WORKING PRINCIPLE
# ══════════════════════════════════════════════════════════════════════
ap('')
ah("Innovation 1: Temporal-Decay PageRank (TD-PageRank)", size=10)
ap("The standard PageRank power iteration computes r[v] = (1-d)/N + d * SUM(r[u]*w(u,v)/SUM(w(u,k))). "
   "In this invention, edge weight w(u,v) is replaced by the temporally-decayed weight:", size=10)
ap("    w_temporal(e) = w_original(e) x exp(-lambda x EdgeAge(e)) x burst_multiplier(sender)",
   bold=True, indent=True, size=10)
ap("where: lambda = 0.693 / half_life_days (default 7 days); EdgeAge = (reference_date - edge.Date).days; "
   "burst_multiplier = 1 + (window_count/total_count) if sender sent >5 transactions in last 3 days, "
   "else 1.0. After convergence (L1 norm < 1e-6, max 100 iterations): asymmetric SCC penalty is applied "
   "(collector 0.25x, distributor 0.50x, balanced 0.375x) and dormant nodes are capped at 0.1x max score.", size=10)
sep()

ah("Innovation 2: Topology-Adaptive Fusion with TopologyAttentionGate", size=10)
ap("For each account, the 2-hop ego-network is extracted and a 6-dimensional TopologyVector is computed. "
   "This vector is passed to the TopologyAttentionGate:", size=10)
ap("    gate       = sigmoid(W_gate x topology_vector)        [values in (0,1)]", indent=True, size=9.5)
ap("    gate_score = mean(gate)                                [topology activity scalar]", indent=True, size=9.5)
ap("    gate_amp   = 1.0 + 2.0 x gate_score                  [amplifier in (1.0, 3.0)]", indent=True, size=9.5)
ap("    base_w     = softmax(W_base x topology_vector)        [base weight allocation]", indent=True, size=9.5)
ap("    w_graph   *= gate_amp                                 [amplify graph weight only]", indent=True, size=9.5)
ap("    w_final    = clamp(w/sum(w), 0.05, 0.90)             [normalize and clamp x3]", indent=True, size=9.5)
ap("    fused      = w_ml*S_ml + w_graph*S_graph + w_rules*S_rules", indent=True, size=9.5)
sep()

ah("Innovation 3: Adaptive Degradation Controller", size=10)
ap("Every 10 seconds, 3-metric HealthVectors are computed for each component. State transitions: "
   "HEALTHY -> DEGRADED after 2 consecutive unhealthy cycles; DEGRADED -> HEALTHY after 3 "
   "consecutive healthy cycles; flap protection locks COOLDOWN for 5 minutes after >3 transitions "
   "in 5 minutes. On any state change: select_path(healthy_set) = argmax P@50 such that "
   "required_components is a subset of healthy_set AND P@50 >= AdaptiveBudget.current(). "
   "Shadow evaluation verifies live precision; PrecisionDriftDetector monitors KL-divergence.", size=10)
sep()

ah("Innovation 4: Architecture Hardening Layer", size=10)
ap("ArchitectureHardeningConfig controls all flags (all default = False — identical to baseline). "
   "When enabled: LearnableSCCPenalty MLP replaces static 3-archetype penalty via train_penalty() "
   "+ ResourceManager.atomic_model_swap(); TopologyEmbeddingNetwork (2x GCNConv, 16-d) + "
   "IsolationDetector replace manual TopologyVector with 200ms timeout/fallback; "
   "DeepTopologyAttentionGate (>=2 hidden layers, >=64 hidden dim, dropout) replaces single-layer "
   "gate; OnlinePrecisionMonitor + EnhancedConceptDriftDetector + EnhancedAdaptivePrecisionBudget "
   "(0.7x online + 0.3x shadow) replace static assumptions; ResourceManager logs every path switch "
   "with memory delta, computation time, and active-unit count.", size=10)
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 7 — DESCRIPTION IN DETAIL
# ══════════════════════════════════════════════════════════════════════
ap('')
ah("7.1  System Architecture", size=10)
aimg("fig01_system_architecture.png", width=6.4,
     cap="Fig 1 — Full system architecture (7 rows): Data Inputs | Feature Eng + LightGBM | "
         "TD-PageRank | Ego-Network | TopologyAttentionGate | Adaptive Fusion | Symbolic Rules | "
         "Degradation Controller | Architecture Hardening Strip | NLP + Dashboard")
ap("The platform processes raw CSV data through seven architectural layers. The core scoring "
   "pipeline flows from Feature Engineering and LightGBM (S_ml), through TD-PageRank (S_graph) "
   "and Topology-Adaptive Fusion, to the Degradation Controller routing through one of eight "
   "pre-validated execution paths. The Architecture Hardening Layer (Row F) is a dedicated "
   "full-width strip, activated only via ArchitectureHardeningConfig flags. NLP pipeline "
   "(spaCy + SmolLM-135M-Instruct) generates entity-faithful STR narratives.", size=10)
sep()

ah("7.2  TD-PageRank Engine  (models/td_pagerank.py)", size=10)
aimg("fig02_tdpagerank_flowchart.png", width=5.5,
     cap="Fig 2 — TD-PageRank algorithm flowchart: temporal weighting → burst amplification → "
         "power iteration → directional SCC penalty → dormant-node suppression")
ap("Class: TDPageRankEngine. Parameters: half_life_days=7.0, damping=0.85, cycle_penalty=0.5, "
   "max_iter=100, tol=1e-6, burst_window_days=3, log_amount_scale=False, "
   "asymmetric_scc_penalty=False, use_learnable_penalty=False.", size=10)
sep()
ap("Full power iteration algorithm:", bold=True, size=10)
ap("  Initialize: r[v] = 1/N for all nodes v", indent=True, size=9.5)
ap("  For t = 1 to max_iter:", indent=True, size=9.5)
ap("    dangling = damping * SUM(r[u] for dangling u) / N", indent=True, size=9.5)
ap("    r_new[v] = (1-d)/N + dangling + d * SUM(r[u]*w_temporal(u,v)/SUM(w_temporal(u,k)))", indent=True, size=9.5)
ap("    if L1(r_new - r) < 1e-6: CONVERGED; break", indent=True, size=9.5)
ap("  Post-processing:", indent=True, size=9.5)
ap("    For each cycle node (SCC size>2, intra_ratio>0.80):", indent=True, size=9.5)
ap("      Collector (inflow-dominant): r[node] *= 0.25", indent=True, size=9.5)
ap("      Distributor (outflow-dominant): r[node] *= 0.50", indent=True, size=9.5)
ap("      Balanced: r[node] *= 0.375", indent=True, size=9.5)
ap("    Dormant nodes (all edges >30 days): r[node] = min(r[node], 0.1 * max(r))", indent=True, size=9.5)
ap("    r_normalized = (r - min(r)) / (max(r) - min(r))", indent=True, size=9.5)
ap("Output: TDPageRankResult — scores, normalized_scores, cycle_member, decay_impact, converged, "
   "iterations, reference_date.", size=10)
sep()

ah("7.3  Topology-Adaptive Fusion  (models/adaptive_fusion.py)", size=10)
aimg("fig03_topology_attention_gate.png", width=6.2,
     cap="Fig 3 — TopologyAttentionGate neural architecture: multiplicative gating on the weight "
         "dimension (not feature-space attention)")
ap("EgoNetworkExtractor computes 6-dimensional TopologyVector from each account's 2-hop ego-network:", size=10)
ab("edge_density = |E| / (|V| x (|V|-1))")
ab("diameter = longest shortest path in undirected ego-network")
ab("avg_clustering = mean clustering coefficient of all ego-network nodes")
ab("degree_asymmetry = Var(in_degree) / max(Var(out_degree), epsilon)")
ab("component_ratio = |largest weakly connected component| / |V|")
ab("tx_velocity_ratio = recent_edges (<=7 days) / total_edges  [optional 6th dimension]")
sep()
ap("TopologyAttentionGate Architecture:", bold=True, size=10)
ab("Layer 1: Linear(5->5) + Sigmoid  [gate layer: per-dimension activity in (0,1)]")
ab("Layer 2: Linear(5->32) + ReLU + Linear(32->3)  [base weight allocator]")
ab("gate_activity = mean(sigmoid_output)  [scalar topology activity]")
ab("gate_amplifier = 1.0 + 2.0 x gate_activity  [in (1.0, 3.0)]")
ab("base_weights = softmax(base_allocator output)  [sums to 1]")
ab("w_graph *= gate_amplifier  [amplify graph signal only]")
ab("weights = clamp(w/sum(w), 0.05, 0.90)  [3 iterations of clamp-normalize]")
sep()
ap("Fallback Logic: Static weights (0.70, 0.15, 0.15) applied when: graph unavailable, topology "
   "extraction >200ms, network inference >50ms, NaN/Inf in outputs, weight outside [0.05, 0.90], "
   "or ego-network isolation (<3 nodes => w_ml forced >=0.70).", size=10)
sep()

ah("7.4  Adaptive Degradation Controller  (models/degradation_controller.py)", size=10)
aimg("fig04_degradation_state_machine.png", width=6.4,
     cap="Fig 4 — State machine (HEALTHY/DEGRADED/COOLDOWN) with transitions (left) and "
         "8-path routing table with P@50 values (right)")
sep()
ap("Routing Table (pre-computed offline):", bold=True, size=10)
mk_table(
    ["Path ID", "Required Components", "P@50", "Level"],
    [
        ["full", "lightgbm, td_pagerank, fusion_engine, symbolic_rules, nlp_summarizer", "0.82", "0"],
        ["no_nlp", "lightgbm, td_pagerank, fusion_engine, symbolic_rules", "0.82", "1"],
        ["no_symbolic", "lightgbm, td_pagerank, fusion_engine", "0.76", "1"],
        ["no_pagerank", "lightgbm, fusion_engine, symbolic_rules", "0.72", "1"],
        ["no_fusion", "lightgbm, td_pagerank, symbolic_rules", "0.70", "1"],
        ["lgbm_pagerank", "lightgbm, td_pagerank", "0.68", "2"],
        ["lgbm_rules", "lightgbm, symbolic_rules", "0.66", "2"],
        ["lgbm_only", "lightgbm", "0.62", "5"],
    ]
)
sep()
ap("PrecisionDriftDetector: EMA-smoothed 20-bin histogram reference (alpha=0.1, Laplace smoothing); "
   "KL-divergence alert when >0.15. AdaptivePrecisionBudget: tightens +0.05 when P@50 consistently "
   ">base+0.10 for 5 estimates; relaxes -0.02 when within 0.05 of base for 3 estimates; "
   "output always in [0.55, 0.75].", size=10)
sep()

ah("7.5  Learnable SCC Penalty  (LearnableSCCPenalty in models/td_pagerank.py)", size=10)
ap("A 2-layer MLP (input_dim=5, hidden_dim=32) maps 5 intra-SCC flow features "
   "[intra_inflow, intra_outflow, weight_ratio, scc_size, node_degree] to penalty multipliers "
   "via sigmoid scaling: penalty = 0.1 + 0.9 * sigma(MLP(x)), guaranteeing output in [0.1, 1.0] "
   "for any input. Training enforces mandatory temporal splits (train_penalty() splits at "
   "temporal_split_date). Config flag: use_learnable_scc_penalty=False (default).", size=10)
sep()

ah("7.6  GNN Topology Embedding  (TopologyEmbeddingNetwork in models/adaptive_fusion.py)", size=10)
ap("Two GCNConv message-passing layers (hidden_dim=32) followed by global mean pooling and a "
   "linear projection to 16-dimensional embedding. Node features: 4-dimensional "
   "[in_degree, out_degree, amount_sum, edge_count]. Timeout enforced at 200ms; fallback to "
   "manual TopologyVector on timeout or PyG unavailability. IsolationDetector: 3-layer MLP "
   "-> sigmoid in [0,1]; compute_fusion_weights() applies: "
   "w_ml_new = w_ml + isolation * (0.90 - w_ml) (monotonically non-decreasing, capped at 0.90).", size=10)
sep()

ah("7.7  DeepTopologyAttentionGate  (models/adaptive_fusion.py)", size=10)
ap("MLP with >=2 hidden layers (hidden_dim>=64) + Dropout + final linear(3) + softmax. "
   "Topology activity = sigmoid(||embedding||_2); graph-signal weight amplified by "
   "1 + 2.0 * activity before iterative clamp-normalize. Inference budget: <50ms. "
   "Raises ValueError if num_hidden_layers<2 or hidden_dim<64.", size=10)
sep()

ah("7.8  Online Precision Recalibration  (models/degradation_controller.py)", size=10)
ap("OnlinePrecisionMonitor: rolling window (window_size=500, min_samples=200); "
   "ingest_feedback() rejects out-of-order timestamps; estimate_precision() returns P@50 "
   "from top-50 scored accounts in window (or None when <200 samples).", size=10)
ap("EnhancedConceptDriftDetector: per-path PrecisionDriftDetector instances (no cross-path "
   "contamination); classify_drift() returns: precision_degraded (precision < budget, regardless "
   "of KL), benign_shift (KL > 0.15 but precision >= budget), no_drift (otherwise).", size=10)
ap("EnhancedAdaptivePrecisionBudget: blended = 0.7*online + 0.3*shadow; tightens +0.05 after "
   "5 consecutive > base+0.10; relaxes -0.02 after 3 consecutive within 0.05 of base; "
   "output clamped to [0.55, 0.75].", size=10)
sep()

ah("7.9  ResourceManager  (models/resource_manager.py)", size=10)
ab("allocate_penalty_buffers(mode): Allocates distinct memory buffers — symmetric: 2 MB, "
   "learnable: 8 MB — for each SCC penalty mode. Req 7.1.")
ab("route_topology_pipeline(use_gnn): Routes to GNN pipeline (80 MB, 120ms, 4 processing units) "
   "or manual pipeline (2 MB, 5ms, 1 unit) with measurably different resource profiles. Req 7.2.")
ab("release_degraded_resources(disabled): Frees component memory on degradation events, "
   "removes disabled components from active set. Req 7.4.")
ab("atomic_model_swap(new_weights): Thread-safe atomic weight replacement (deep-copy + lock) "
   "without interrupting scoring; rollback on failure. Req 7.6.")
ab("utilization_report(): Generates ResourceUtilizationSnapshot records with memory delta, "
   "computation time, and active-unit count per event. Req 7.3, 7.5.")
sep()

ah("7.10  System Data Flow", size=10)
aimg("fig05_data_flow.png", width=6.4,
     cap="Fig 5 — End-to-end data flow: raw CSVs -> feature engineering -> TD-PageRank / "
         "LightGBM / symbolic rules -> topology-adaptive fusion -> degradation controller -> "
         "account aggregation -> analyst dashboard")
sep()
aimg("fig09_conceptual_innovations.png", width=6.4,
     cap="Fig 9 — Conceptual overview: (left) TD-PageRank on transaction network; "
         "(centre) per-account fusion weights for isolated vs hub accounts; "
         "(right) degradation controller maintaining P@50 across component failures")
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 8 — EXPERIMENTAL VALIDATION
# ══════════════════════════════════════════════════════════════════════
ap('')
ah("8.1  System Performance on Temporal Test Set", size=10)
ap("All results measured on the temporal test set (chronologically last 17% of data, Oct-Dec). "
   "No future data leakage. random_state=42. NPR-denominated Nepali banking transaction dataset.", size=10)
sep()
mk_table(
    ["Metric", "Full System", "Static Fusion Baseline", "Standard PageRank", "Random Sampling"],
    [
        ["Precision@10",  "0.92", "0.80", "0.75", "0.005"],
        ["Precision@50",  "0.82", "0.62", "0.55", "0.005"],
        ["AUC-PR",        "0.63", "0.52", "0.46", "0.005"],
        ["NDCG@50",       "0.48", "0.35", "0.30", "~0"],
        ["Lift@50",       "~98x", "~68x", "~55x", "1x"],
        ["Brier Score",   "0.009","0.018","0.024","0.005"],
    ]
)
sep()

ah("8.2  Ablation Study — Incremental Component Value", size=10)
aimg("fig06_ablation_study.png", width=6.0,
     cap="Fig 6 — Ablation study: incremental P@50 improvement from each system component")
mk_table(
    ["Configuration", "Precision@50", "Increment"],
    [
        ["Tabular LightGBM only",                    "0.41", "Baseline"],
        ["+ Graph Features (16 centrality metrics)", "0.55", "+0.14"],
        ["+ Symbolic Rules (10 FATF typologies)",    "0.60", "+0.05"],
        ["+ Static Fusion (fixed weights)",          "0.65", "+0.05"],
        ["+ TD-PageRank + TopologyAttentionGate",    "0.76", "+0.11"],
        ["+ Degradation Controller (full system)",   "0.82", "+0.06"],
    ]
)
sep()

ah("8.3  Prior-Art Differentiation Thresholds — Measured", size=10)
mk_table(
    ["Innovation", "vs Prior Art", "Metric", "Required", "Measured", "Status"],
    [
        ["TD-PageRank", "US11640609B1", "Mean Absolute Score Difference from standard PageRank", ">=0.01 (1%)", "0.027 (2.7%)", "PASS"],
        ["TD-PageRank", "US11640609B1", "Mean Absolute Percentage Difference", ">=15%", "18.3%", "PASS"],
        ["Fusion", "US20220405860A1", "Relative P@50 improvement over static fusion", ">=10%", "32.3%", "PASS"],
        ["Degradation", "US20210174258A1", "Min P@50 across all 8 execution paths", ">=0.60", "0.62 (lgbm_only)", "PASS"],
        ["Degradation", "US20210174258A1", "Self-calibrating precision budget", "Dynamic [0.55-0.75]", "AdaptiveBudget.current()", "DIFFERENTIATOR"],
    ]
)
sep()

ah("8.4  Routing Table Precision Validation", size=10)
aimg("fig08_precision_routing_table.png", width=6.0,
     cap="Fig 8 — All 8 execution paths meet P@50 >= 0.60 precision budget guarantee")
sep()

ah("8.5  Property-Based Test Results (119 Tests, All PASS)", size=10)
aimg("fig10_property_test_coverage.png", width=6.4,
     cap="Fig 10 — 27 formal correctness properties verified by Hypothesis framework (50-100 examples each). All 119 pytest tests pass.")
sep()

ah("8.6  Architecture Hardening Property Summary", size=10)
mk_table(
    ["ID", "Validates", "Formal Statement", "Status"],
    [
        ["AH-1", "Req 1.3", "LearnableSCCPenalty output in [0.1,1.0] for any valid SCCFlowFeatures", "PASS"],
        ["AH-2", "Req 1.7", "Two consecutive calls produce identical output within 1e-12", "PASS"],
        ["AH-3", "Req 1.5", "use_learnable_penalty=False reproduces static heuristic exactly", "PASS"],
        ["AH-4", "Req 2.1", "TopologyEmbeddingNetwork always outputs exactly embedding_dim dimensions", "PASS"],
        ["AH-5", "Req 2.6,3.4,3.5", "<3 nodes => isolation>=0.9; >=10 nodes density>0.3 => isolation<0.2", "PASS"],
        ["AH-6", "Req 3.1,3.3,3.7", "Isolation weights: each in [0.05,0.90], sum=1.0+-1e-6, w_ml monotone", "PASS"],
        ["AH-7", "Req 4.1,4.3,4.5", "Monitor: None when <200 samples; valid P@50 when >=200; timestamps monotone", "PASS"],
        ["AH-8", "Req 4.2", "Routing table updated iff |p_live - p_stored| > 0.05", "PASS"],
        ["AH-9", "Req 5.2,5.3,5.4", "Drift classification correctness + per-path isolation (no cross-contamination)", "PASS"],
        ["AH-10", "Req 6.2,6.3", "Deep gate: each weight in [0.05,0.90], sum=1.0+-1e-6, multiplicative property", "PASS"],
        ["AH-11", "Req 8.2-8.5", "Budget: 0.7*online+0.3*shadow; tighten/relax logic; output in [0.55,0.75]", "PASS"],
        ["AH-12", "Req 7.3", "Every path switch produces log with non-zero memory delta, time>=0, units>0", "PASS"],
    ]
)
sep()

ah("8.7  Leakage Verification", size=10)
ap("Running with random (non-temporal) train/test split artificially inflates P@50 to 0.97. "
   "Our temporal split gives 0.82. The 0.15 difference quantifies the exact magnitude of leakage "
   "that would have been present without temporal discipline. The reported 0.82 represents "
   "genuine generalization on unseen future transactions.", size=10)
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 9 — WHAT NEEDS PROTECTION
# ══════════════════════════════════════════════════════════════════════
ap('')
ah("Claim 1 — Temporal-Decay PageRank with Burst-Velocity Amplification "
   "and Directional SCC Penalty", size=10)
ap("What must be protected:", bold=True, size=10)
reset_counter()
an("1.1  The method of embedding exp(-lambda x EdgeAge) directly within the PageRank power "
   "iteration transition matrix (not as pre-processing on time-binned snapshots). This is the "
   "core algorithmic novelty.")
an("1.2  Burst-velocity amplification: per-edge weight multiplier (1 + window_count/total_count) "
   ">=1.0 applied to sender nodes exhibiting transaction surges within a configurable time window.")
an("1.3  Log-amount normalisation: replacing raw transaction amounts with log(1+amount) before "
   "temporal weighting to prevent large single transactions from dominating centrality.")
an("1.4  Directional SCC penalty: asymmetric dampening distinguishing fund collectors "
   "(inflow-dominant, 0.25x), distributors (outflow-dominant, 0.50x), and balanced nodes (0.375x).")
an("1.5  Dormant node suppression: capping rank scores at 0.1x the maximum when all incident "
   "edges exceed a configurable age threshold (default 30 days).")
an("1.6  (New) LearnableSCCPenalty MLP: trainable replacement for static 3-archetype penalty, "
   "retrained without code changes, updated atomically via ResourceManager.atomic_model_swap().")
ap("Scope: Method claims for the combination (1.1)-(1.6) applied to directed financial transaction "
   "graphs with timestamped, amount-weighted edges.", italic=True, size=9.5)
sep()

ah("Claim 2 — Topology-Conditioned Multiplicative Attention Gate "
   "for Heterogeneous Signal Fusion", size=10)
ap("What must be protected:", bold=True, size=10)
reset_counter()
an("2.1  Per-account 5- or 6-dimensional TopologyVector from 2-hop ego-network.")
an("2.2  TopologyAttentionGate: multiplicative gating on the ensemble weight dimension (NOT "
   "additive attention on feature space); gate amplifier in [1.0, 3.0].")
an("2.3  Complete per-account weight pipeline: TopologyVector -> gate -> clamped weights [0.05, 0.90] "
   "-> fused score; computed independently per account per batch.")
an("2.4  Isolation safety floor: w_ml >=0.70 when ego-network <3 nodes.")
an("2.5  Complete 6-condition fallback chain to static weights (0.70, 0.15, 0.15).")
an("2.6  (New) TopologyEmbeddingNetwork: GNN-based learned ego-network embedding "
   "(2x GCNConv, 16-d, 200ms timeout with fallback).")
an("2.7  (New) IsolationDetector: learned continuous isolation function, "
   "monotonically non-decreasing w_ml, replaces hard-coded threshold.")
an("2.8  (New) DeepTopologyAttentionGate: >=2 hidden layers, dropout, deeper multiplicative gating.")
ap("Scope: System and method claims covering (2.1)-(2.8) as per-account dynamic ensemble weighting "
   "for heterogeneous AML risk signal fusion.", italic=True, size=9.5)
sep()

ah("Claim 3 — Precision-Budget-Constrained Adaptive Degradation Controller", size=10)
ap("What must be protected:", bold=True, size=10)
reset_counter()
an("3.1  Pre-computed execution path routing table with P@K annotations; selection constrained "
   "by precision budget (not latency/availability metrics).")
an("3.2  PrecisionDriftDetector: EMA-smoothed KL-divergence for proactive precision degradation "
   "detection before it manifests in measured P@K values.")
an("3.3  AdaptivePrecisionBudget: self-calibrating minimum P@K threshold with rolling shadow "
   "evaluation history; tightens/relaxes; bounded by configurable floor and ceiling.")
an("3.4  Health-monitoring state machine: 2-cycle DEGRADED detection, 3-cycle HEALTHY recovery, "
   "flap protection, 30-min stale queue policy, <=500ms routing SLA.")
an("3.5  Drift KL-divergence as additional input to budget computation (+0.03 tightening).")
an("3.6  (New) OnlinePrecisionMonitor: live P@50 from rolling labeled feedback "
   "(window=500, min=200 samples); timestamp ordering validation.")
an("3.7  (New) EnhancedConceptDriftDetector: KL + precision jointly classify no_drift / "
   "benign_shift / precision_degraded with per-path isolation.")
an("3.8  (New) EnhancedAdaptivePrecisionBudget: 0.7x online + 0.3x shadow blending with "
   "consecutive-estimate tightening/relaxation logic.")
ap("Scope: System and method claims covering (3.1)-(3.8) as output-quality-preserving adaptive "
   "degradation controller for multi-component ML scoring pipelines.", italic=True, size=9.5)
sep()

ah("Claim 4 — Patent Evaluation Harness (Supporting Claim)", size=10)
reset_counter()
an("4.1  Reproducible evaluation harness measuring each innovation against identified prior art "
   "using fixed random seeds, chronological temporal splits, and automatic flagging when "
   "differentiation thresholds are not met.")
ap("Scope: Method claims for the evaluation harness as a reproducible prior-art differentiation "
   "measurement system.", italic=True, size=9.5)
sep()

ah("Claim 5 — §101 Resource Management Anchoring (ResourceManager)", size=10)
ap("What must be protected:", bold=True, size=10)
reset_counter()
an("5.1  allocate_penalty_buffers(mode): Distinct memory buffers (symmetric: 2MB; learnable: 8MB) "
   "for each SCC penalty mode with measurably different sizes.")
an("5.2  route_topology_pipeline(use_gnn): Routes to GNN pipeline (80MB, 120ms, 4 units) or "
   "manual pipeline (2MB, 5ms, 1 unit) with measurably different memory and processing profiles.")
an("5.3  release_degraded_resources(disabled): Frees component memory on degradation; "
   "demonstrates reduced memory usage during degraded operation.")
an("5.4  atomic_model_swap(new_weights): Thread-safe atomic weight update without interrupting "
   "scoring; rollback on failure.")
an("5.5  utilization_report(): Per-event ResourceUtilizationSnapshot (memory delta, computation "
   "time, active-unit count) providing concrete §101 technological improvement evidence.")
ap("Scope: System and method claims for (5.1)-(5.5) as a concrete resource management layer "
   "tied to algorithmic innovations, surviving Alice Corp. v. CLS Bank abstract-idea challenges.",
   italic=True, size=9.5)
sep()

# ══════════════════════════════════════════════════════════════════════
# SECTION 10 — TRL
# ══════════════════════════════════════════════════════════════════════
ap('')
ap("Selected TRL:  TRL 6 — Technology demonstrated in a relevant environment",
   bold=True, size=11)
sep()
mk_table(
    ["Stage", "TRL", "Description", "Status for This Invention"],
    [
        ["Research",    "TRL 1", "Basic principles observed",
         "COMPLETE — All three innovations founded on proven mathematical principles "
         "(PageRank, attention gating, circuit-breaker patterns)"],
        ["Research",    "TRL 2", "Technology concept formulated",
         "COMPLETE — Formal problem statements and novelty gaps documented against 10 prior art references"],
        ["Research",    "TRL 3", "Experimental proof of concept",
         "COMPLETE — All innovations implemented in Python; 119 property-based tests prove correctness of "
         "27 mathematical invariants (original 15 + 12 architecture-hardening)"],
        ["Development", "TRL 4", "Technology validated in a lab",
         "COMPLETE — Full ablation study confirms P@50: 0.41 (tabular only) -> 0.82 (full system); "
         "architecture hardening adds 12 new verified properties"],
        ["Development", "TRL 5", "Technology validated in relevant environment",
         "COMPLETE — Validated on real-world NPR-denominated transaction data; temporal split prevents "
         "leakage; all 3 prior-art differentiation thresholds met (MAD 2.7%, P@50 gain 32.3%, min P@50 0.62)"],
        ["Development", "TRL 6", "Technology demonstrated in relevant environment",
         "COMPLETE — Full production-deployable Streamlit dashboard (streamlit run app.py); "
         "Degradation Controller tested under simulated component failures; Architecture Hardening "
         "layer operational with ArchitectureHardeningConfig; ResourceManager generating §101 evidence; "
         "119 tests all passing"],
        ["Deployment",  "TRL 7", "System prototype in operational environment",
         "In Progress — Platform architecturally complete; live bank transaction feed integration pending"],
        ["Deployment",  "TRL 8", "System complete and qualified",
         "Not yet — Requires institutional deployment and security audit"],
        ["Deployment",  "TRL 9", "Actual system proven in operational environment",
         "Not yet — Requires production deployment at a financial institution"],
    ]
)
sep()
ap("Justification for TRL 6:", bold=True, size=10)
ap("The full platform runs end-to-end from raw CSV input to prioritised analyst queue with one "
   "command (streamlit run app.py). All three original innovations and all four architecture-hardening "
   "surfaces have been implemented, tested (119 passing tests across 27 formal properties), and "
   "validated on realistic NPR-denominated financial transaction data. The Degradation Controller "
   "operates under simulated component failures maintaining P@50 >=0.60 across all 8 execution paths. "
   "The ArchitectureHardeningConfig layer enables seamless backward-compatible activation of all "
   "new components. The ResourceManager generates concrete §101 evidence (memory delta, computation "
   "time, active-unit count) at runtime. The platform is ready for pilot deployment at a financial "
   "institution pending data-feed integration and security audit.", size=10)
sep()

# ══════════════════════════════════════════════════════════════════════
# Restore END paragraph + save
# ══════════════════════════════════════════════════════════════════════
if end_para_el is not None:
    doc.element.body.append(end_para_el)
else:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("----------------------END OF THE DOCUMENT-----------------------------")
    r.bold = True
    r.font.size = Pt(10)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Document No.: 02-IPR-R003  |  Issue No/Date: 2/01.02.2024  |  Amendment No/Date: 1/22.06.2026")
    r2.italic = True
    r2.font.size = Pt(9)

doc.save(OUT)
print(f"\nSaved: {OUT}")
print(f"Size: {os.path.getsize(OUT)/1024:.1f} KB")
